import os
import re
import zipfile
import copy
import tempfile
from pathlib import Path
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================================
# 1. XML UTILITIES (GIỮ NGUYÊN HOÀN TOÀN TỪ BẢN TRƯỚC)
# ==============================================================================

def copy_xml_element(src_element, dest_element, prop_tag: str):
    if src_element is None or dest_element is None: return
    old_pr = dest_element.find(prop_tag, namespaces=dest_element.nsmap)
    if old_pr is not None: dest_element.remove(old_pr)
    src_pr = src_element.find(prop_tag, namespaces=src_element.nsmap)
    if src_pr is not None: dest_element.insert(0, copy.deepcopy(src_pr))

def apply_visual_paragraph_properties(raw_pPr, tpl_pPr):
    if tpl_pPr is None: return copy.deepcopy(raw_pPr) if raw_pPr is not None else None
    merged_pPr = copy.deepcopy(tpl_pPr)
    namespaces = merged_pPr.nsmap
    tpl_num = merged_pPr.find(qn('w:numPr'))
    if tpl_num is not None:
        merged_pPr.remove(tpl_num)
        tpl_ind = merged_pPr.find(qn('w:ind'))
        if tpl_ind is not None: merged_pPr.remove(tpl_ind)
    if raw_pPr is not None:
        raw_num = raw_pPr.find(qn('w:numPr'))
        if raw_num is not None:
            merged_pPr.append(copy.deepcopy(raw_num))
            raw_ind = raw_pPr.find(qn('w:ind'))
            if raw_ind is not None:
                old_ind = merged_pPr.find(qn('w:ind'))
                if old_ind is not None: merged_pPr.remove(old_ind)
                merged_pPr.append(copy.deepcopy(raw_ind))
    return merged_pPr

def apply_visual_run_properties(raw_rPr, tpl_rPr, preserve_char: bool = True):
    if tpl_rPr is None: return raw_rPr
    merged_rPr = copy.deepcopy(tpl_rPr)
    if raw_rPr is not None and preserve_char:
        tags_to_keep = [qn('w:b'), qn('w:bCs'), qn('w:i'), qn('w:iCs'), qn('w:u'), qn('w:highlight')]
        for tag in tags_to_keep:
            raw_tag = raw_rPr.find(tag)
            if raw_tag is not None:
                old_tag = merged_rPr.find(tag)
                if old_tag is not None: merged_rPr.remove(old_tag)
                merged_rPr.append(copy.deepcopy(raw_tag))
    return merged_rPr

def cleanse_raw_run(rPr):
    if rPr is None: return
    for child in list(rPr):
        if child.tag.endswith('}rFonts') or child.tag.endswith('}sz') or child.tag.endswith('}szCs') or child.tag.endswith('}color'):
            rPr.remove(child)

def force_zero_cell_spacing(tblPr):
    if tblPr is None: return
    for tag in tblPr.findall(qn('w:tblCellSpacing')): tblPr.remove(tag)
    zero_spacing_xml = '<w:tblCellSpacing w:w="0" w:type="nil" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    nil_spacing = parse_xml(zero_spacing_xml)
    insert_idx = 0
    before_tags = [qn('w:tblStyle'), qn('w:tblpPr'), qn('w:tblOverlap'), qn('w:bidiVisual'), qn('w:tblStyleRowBandSize'), qn('w:tblStyleColBandSize'), qn('w:tblW'), qn('w:jc')]
    for i, child in enumerate(tblPr):
        if child.tag in before_tags: insert_idx = max(insert_idx, i + 1)
    tblPr.insert(insert_idx, nil_spacing)

def is_auto_bullet(doc, pPr):
    try:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None: return False
        numId_node = numPr.find(qn('w:numId'))
        if numId_node is None: return False
        num_id = numId_node.get(qn('w:val'))
        num_part = doc.part.numbering_part
        if num_part is None: return True 
        w_nums = num_part._element.findall(qn('w:num'))
        abs_num_id = None
        for w_num in w_nums:
            if w_num.get(qn('w:numId')) == num_id:
                abs_num_node = w_num.find(qn('w:abstractNumId'))
                if abs_num_node is not None: abs_num_id = abs_num_node.get(qn('w:val'))
                break
        if abs_num_id is None: return True
        ilvl_node = numPr.find(qn('w:ilvl'))
        lvl_val = ilvl_node.get(qn('w:val')) if ilvl_node is not None else "0"
        w_abs_nums = num_part._element.findall(qn('w:abstractNum'))
        for w_abs in w_abs_nums:
            if w_abs.get(qn('w:abstractNumId')) == abs_num_id:
                for w_lvl in w_abs.findall(qn('w:lvl')):
                    if w_lvl.get(qn('w:ilvl')) == lvl_val:
                        numFmt = w_lvl.find(qn('w:numFmt'))
                        if numFmt is not None: return numFmt.get(qn('w:val')) == 'bullet'
        return True
    except Exception:
        return True

# ==============================================================================
# 2. VISUAL TEMPLATE ENGINE 
# ==============================================================================

class VisualTemplateEngine:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = None
        self.visual_formats = {}
        self.table_format = None
        self.sections = []

    def analyze(self):
        try:
            self.doc = Document(self.template_path)
            self._extract_visual_anchors()
            self._extract_table_anchors()
            self.sections = self.doc.sections
            return True, "Success"
        except Exception as e:
            return False, f"Lỗi phân tích file mẫu: {e}"

    def _extract_visual_anchors(self):
        fallback_pPr, fallback_rPr = None, None
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text: continue
            cat = self.identify_category(text)
            pPr = p._p.find(qn('w:pPr'))
            rPr = None
            if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
                rPr = p.runs[0]._r.find(qn('w:rPr'))
            is_numbered = (pPr is not None and pPr.find(qn('w:numPr')) is not None)
            if cat == "Normal":
                if not is_numbered and fallback_pPr is None:
                    fallback_pPr, fallback_rPr = copy.deepcopy(pPr), copy.deepcopy(rPr)
            else:
                if cat not in self.visual_formats:
                    self.visual_formats[cat] = (copy.deepcopy(pPr), copy.deepcopy(rPr))
        if fallback_pPr is None: fallback_pPr = OxmlElement('w:pPr')
        if fallback_rPr is not None:
            for tag in [qn('w:b'), qn('w:bCs'), qn('w:i'), qn('w:iCs')]:
                elem = fallback_rPr.find(tag)
                if elem is not None: fallback_rPr.remove(elem)
        for key in ["UNIT", "PART_MAIN", "PART_SUB", "EXERCISE", "HEADING_UNNUMBERED", "SUB_ALPHA", "LIST", "Normal", "BAND_HEADING"]:
            if key not in self.visual_formats: self.visual_formats[key] = (fallback_pPr, fallback_rPr)

    def _extract_table_anchors(self):
        if not self.doc.tables: return
        tbl = self.doc.tables[0]
        tblPr = copy.deepcopy(tbl._tbl.find(qn('w:tblPr')))
        row_tcPr_list = []
        if tbl.rows:
            for cell in tbl.rows[0].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                row_tcPr_list.append(copy.deepcopy(tcPr))
        self.table_format = (tblPr, row_tcPr_list)

    def identify_category(self, text: str) -> str:
        clean = text.strip()
        if not clean: return "Normal"
        lower = clean.lower()
        if re.match(r"^(UNIT|Unit)\s+\d+", clean): return "UNIT"
        if re.match(r"^(PART|Part)\s+\d+\.", clean): return "PART_MAIN"
        if re.match(r"^(PART|Part)\s+\d+$", clean): return "PART_SUB"
        if re.match(r"^(Exercise|Task)\s+\d+", clean, re.IGNORECASE): return "EXERCISE"
        if (lower.startswith("reading comprehension") or lower.startswith("sample essay") or lower.startswith("useful vocabulary") or lower.startswith("useful sentence starters") or "band builder" in lower): return "HEADING_UNNUMBERED"
        if re.match(r"^Band\s+\d", clean, re.IGNORECASE): return "BAND_HEADING"
        if re.match(r"^[A-Z]\.\s+[A-Za-z0-9]", clean): return "SUB_ALPHA"
        if re.match(r"^[a-d]\)\s+", clean, re.IGNORECASE) or re.match(r"^\d+\.\s+[A-Za-z0-9]", clean): return "LIST"
        return "Normal"

# ==============================================================================
# 3. WORD FORMATTER PIPELINE 
# ==============================================================================

class WordFormatterPipeline:
    def __init__(self, raw_path, template_engine, output_path, config, progress_callback):
        self.raw_path = raw_path
        self.engine = template_engine
        self.output_path = output_path
        self.config = config
        self.progress = progress_callback
        self.doc = None

    def execute(self):
        try:
            self.progress(10, "Đang nạp tài liệu dữ liệu thô...")
            self.doc = Document(self.raw_path)
            self._preprocess_structure()
            if self.config.get("paragraph", True) or self.config.get("heading", True): self._format_paragraphs()
            if self.config.get("table", True): self._format_tables()
            if self.config.get("page_setup", True): self._format_page_setup()
            self.progress(85, "Đang đồng bộ Styles hệ thống...")
            self.doc.save(self.output_path)
            self._inject_template_styles()
            self.progress(100, "Định dạng thành công!")
            return True, "Hoàn tất định dạng văn bản."
        except Exception as e:
            return False, str(e)

    def _preprocess_structure(self):
        paragraphs = self.doc.paragraphs
        i = 0
        while i < len(paragraphs) - 1:
            p = paragraphs[i]
            text = p.text.strip()
            if re.match(r"^(PART|UNIT)\s+\d+\.$", text, re.IGNORECASE):
                next_p = paragraphs[i+1]
                if next_p.text.strip():  
                    p.text = text + " " + next_p.text.strip()
                    next_element = next_p._element
                    next_element.getparent().remove(next_element)
                    paragraphs = self.doc.paragraphs
                    continue
            i += 1

    def _format_paragraphs(self):
        total_p = len(self.doc.paragraphs)
        current_main_section = ""
        current_sub_section = ""
        numbering_counter = 1
        
        last_text_idx = -1
        for i in range(total_p - 1, -1, -1):
            if self.doc.paragraphs[i].text.strip():
                last_text_idx = i
                break

        for idx, p in enumerate(self.doc.paragraphs):
            if idx % 10 == 0: self.progress(15 + int(40 * (idx / max(total_p, 1))), f"Đang chuẩn hóa đoạn văn {idx}/{total_p}...")
            text = p.text.strip()
            lower_text = text.lower()
            cat = self.engine.identify_category(text)

            if "example sentences" in lower_text or "part 8" in lower_text:
                current_main_section = "EXAMPLE"
                numbering_counter = 1
            elif "speaking" in lower_text or "part 10" in lower_text: current_main_section = "SPEAKING"
            elif "writing" in lower_text or "part 11" in lower_text: current_main_section = "WRITING"
            elif "band builder" in lower_text or "part 13" in lower_text: current_main_section = "BAND_BUILDER"

            if current_main_section == "SPEAKING" and re.match(r"^(part|task)\s+\d+", lower_text):
                current_sub_section = re.match(r"^(part|task)\s+\d+", lower_text).group(0)
                numbering_counter = 1

            raw_pPr = p._p.find(qn('w:pPr'))
            is_bullet = False
            is_manual_numbering = bool(re.match(r"^\d+\.", text))
            
            if cat == "Normal" and not is_manual_numbering:
                if text.startswith("•") or text.startswith("·") or text.startswith("-") or text.startswith("*"): is_bullet = True
                elif raw_pPr is not None and raw_pPr.find(qn('w:numPr')) is not None:
                    if is_auto_bullet(self.doc, raw_pPr): is_bullet = True
                elif current_main_section == "BAND_BUILDER" and text: is_bullet = True
                    
            if idx == last_text_idx:
                is_bullet = True
                is_manual_numbering = False

            tpl_pPr, tpl_rPr = self.engine.visual_formats.get(cat, (None, None))

            if tpl_pPr is not None and cat in ["UNIT", "PART_MAIN", "PART_SUB", "EXERCISE", "HEADING_UNNUMBERED", "SUB_ALPHA", "BAND_HEADING"]:
                tpl_pPr = copy.deepcopy(tpl_pPr)
                spacing = tpl_pPr.find(qn('w:spacing'))
                if spacing is not None: spacing.set(qn('w:before'), '120')
                else:
                    new_sp = OxmlElement('w:spacing')
                    new_sp.set(qn('w:before'), '120')
                    tpl_pPr.append(new_sp)

            if self.config.get("paragraph", True):
                new_pPr = apply_visual_paragraph_properties(raw_pPr, tpl_pPr)
                if raw_pPr is not None: p._p.remove(raw_pPr)
                if new_pPr is not None: p._p.insert(0, new_pPr)
                if cat == "Normal": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if self.config.get("font", True):
                preserve = self.config.get("preserve_char", True)
                for r in p.runs:
                    raw_rPr = r._r.find(qn('w:rPr'))
                    cleanse_raw_run(raw_rPr)
                    new_rPr = apply_visual_run_properties(raw_rPr, tpl_rPr, preserve)
                    if raw_rPr is not None: r._r.remove(raw_rPr)
                    if new_rPr is not None: r._r.insert(0, copy.deepcopy(new_rPr))

            if idx == 0 and text:
                for r in p.runs: r.text = r.text.upper()

            if "a. multiple choice" in lower_text or "b. true / false / not given" in lower_text:
                for r in p.runs:
                    r.bold = False
                    if r._r.rPr is not None:
                        for b in r._r.rPr.findall(qn('w:b')): r._r.rPr.remove(b)
                        for b in r._r.rPr.findall(qn('w:bCs')): r._r.rPr.remove(b)

            if current_main_section == "SPEAKING" and current_sub_section == "part 2" and cat != "PART_SUB":
                for r in p.runs:
                    r.bold = False
                    if r._r.rPr is not None:
                        for b in r._r.rPr.findall(qn('w:b')): r._r.rPr.remove(b)
                        for b in r._r.rPr.findall(qn('w:bCs')): r._r.rPr.remove(b)

            if current_main_section == "WRITING":
                if lower_text in ["useful vocabulary", "useful sentence starters"]:
                    for r in p.runs:
                        r.bold = False
                        if r._r.rPr is not None:
                            for b in r._r.rPr.findall(qn('w:b')): r._r.rPr.remove(b)
                            for b in r._r.rPr.findall(qn('w:bCs')): r._r.rPr.remove(b)
                elif lower_text == "sample essay":
                    for r in p.runs:
                        r.bold = True
                        r.italic = True
                        if r._r.rPr is None: r._r.get_or_add_rPr()
                        if r._r.rPr.find(qn('w:b')) is None: r._r.rPr.append(OxmlElement('w:b'))
                        if r._r.rPr.find(qn('w:i')) is None: r._r.rPr.append(OxmlElement('w:i'))

            if cat == "BAND_HEADING" or (current_main_section == "BAND_BUILDER" and lower_text.startswith("band ")):
                for r in p.runs:
                    r.bold = False
                    if r._r.rPr is not None:
                        for b in r._r.rPr.findall(qn('w:b')): r._r.rPr.remove(b)
                        for b in r._r.rPr.findall(qn('w:bCs')): r._r.rPr.remove(b)

            if is_bullet:
                pPr = p._p.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None: pPr.remove(numPr)
                ind = pPr.find(qn('w:ind'))
                if ind is not None: pPr.remove(ind)
                new_ind = OxmlElement('w:ind')
                new_ind.set(qn('w:left'), '720')
                new_ind.set(qn('w:hanging'), '360')
                pPr.append(new_ind)
                if p.runs: p.runs[0].text = re.sub(r'^[•·\-*]\s*', '', p.runs[0].text.lstrip())
                clean_text = p.text.strip()
                
                is_numbering_context = False
                if current_main_section == "EXAMPLE": is_numbering_context = True
                elif current_main_section == "SPEAKING" and current_sub_section in ["part 1", "part 3"]: is_numbering_context = True
                if idx == last_text_idx: is_numbering_context = False

                if is_numbering_context:
                    if not re.match(r"^\d+\.", clean_text):
                        if p.runs: p.runs[0].text = f"{numbering_counter}. " + p.runs[0].text
                        else: p.add_run(f"{numbering_counter}. ")
                    numbering_counter += 1
                else:
                    if not clean_text.startswith("-"):
                        if p.runs: p.runs[0].text = "- " + p.runs[0].text
                        else: p.add_run("- ")

    def _format_tables(self):
        if not self.engine.table_format: return
        tpl_tblPr, tpl_tcPr_list = self.engine.table_format
        total_tbl = len(self.doc.tables)
        for idx, tbl in enumerate(self.doc.tables):
            self.progress(60 + int(20 * (idx / max(total_tbl, 1))), f"Đang vẽ lại bảng biểu {idx + 1}/{total_tbl}...")
            for pr_ex in tbl._tbl.findall(qn('w:tblPrEx')): tbl._tbl.remove(pr_ex)
            raw_tblPr = tbl._tbl.find(qn('w:tblPr'))
            if raw_tblPr is not None: tbl._tbl.remove(raw_tblPr)
            tpl_tblPr_safe = copy.deepcopy(tpl_tblPr) if tpl_tblPr is not None else parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            force_zero_cell_spacing(tpl_tblPr_safe)
            tbl._tbl.insert(0, tpl_tblPr_safe)

            for row in tbl.rows:
                if row._tr.trPr is not None:
                    for spacing in row._tr.trPr.findall(qn('w:tblCellSpacing')): row._tr.trPr.remove(spacing)
                for c_idx, cell in enumerate(row.cells):
                    tpl_tc = tpl_tcPr_list[c_idx] if c_idx < len(tpl_tcPr_list) else (tpl_tcPr_list[-1] if tpl_tcPr_list else None)
                    raw_tcPr = cell._tc.find(qn('w:tcPr'))
                    if raw_tcPr is not None: cell._tc.remove(raw_tcPr)
                    if tpl_tc is not None: cell._tc.insert(0, copy.deepcopy(tpl_tc))

    def _format_page_setup(self):
        if not self.engine.sections or not self.doc.sections: return
        tpl_sec = self.engine.sections[0]
        for sec in self.doc.sections:
            sec.page_width, sec.page_height = tpl_sec.page_width, tpl_sec.page_height
            for margin in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin', 'header_distance', 'footer_distance']:
                setattr(sec, margin, getattr(tpl_sec, margin))

    def _inject_template_styles(self):
        styles_xml = None
        with zipfile.ZipFile(self.engine.template_path, 'r') as z_tpl:
            if 'word/styles.xml' in z_tpl.namelist(): styles_xml = z_tpl.read('word/styles.xml')
        if not styles_xml: return
        temp_path = self.output_path + '.tmp'
        with zipfile.ZipFile(self.output_path, 'r') as z_in:
            with zipfile.ZipFile(temp_path, 'w', compression=zipfile.ZIP_DEFLATED) as z_out:
                for item in z_in.infolist():
                    if item.filename == 'word/styles.xml': z_out.writestr(item, styles_xml)
                    else: z_out.writestr(item, z_in.read(item.filename))
        os.replace(temp_path, self.output_path)

# ==============================================================================
# 4. STREAMLIT WEB APP UI
# ==============================================================================

st.set_page_config(page_title="Word Formatter Pro", page_icon="📝", layout="centered")

st.title("📝 Word Formatter Pro")
st.markdown("Công cụ tự động đồng bộ và chuẩn hóa định dạng Microsoft Word - By Đặng Văn Nghĩa")

st.markdown("### 1. Upload tài liệu")
col1, col2 = st.columns(2)
with col1:
    template_file = st.file_uploader("Tải lên File Mẫu (.docx)", type=["docx"])
with col2:
    raw_file = st.file_uploader("Tải lên File Dữ Liệu Thô (.docx)", type=["docx"])

st.markdown("### 2. Cấu hình định dạng")
c1, c2 = st.columns(2)
with c1:
    cfg_font = st.checkbox("Font & Cỡ chữ", value=True)
    cfg_paragraph = st.checkbox("Khoảng cách dòng & Đoạn", value=True)
    cfg_heading = st.checkbox("Cấu trúc Heading", value=True)
with c2:
    cfg_table = st.checkbox("Định dạng Bảng biểu (Khít viền)", value=True)
    cfg_page_setup = st.checkbox("Khổ giấy & Căn lề", value=True)
    cfg_preserve_char = st.checkbox("Bảo tồn chữ In Đậm/Nghiêng", value=True)

if st.button("🚀 TIẾN HÀNH ĐỊNH DẠNG", use_container_width=True):
    if not template_file or not raw_file:
        st.error("Vui lòng tải lên đầy đủ File Mẫu và File Dữ Liệu Thô!")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()

        def update_progress(percent, msg):
            progress_bar.progress(percent / 100.0)
            status_text.text(msg)

        # Tạo thư mục tạm để xử lý file
        with tempfile.TemporaryDirectory() as tmpdirname:
            tpl_path = os.path.join(tmpdirname, template_file.name)
            raw_path = os.path.join(tmpdirname, raw_file.name)
            output_name = raw_file.name.replace(".docx", "_Formatted.docx")
            out_path = os.path.join(tmpdirname, output_name)

            # Lưu file upload xuống ổ đĩa cục bộ tạm thời
            with open(tpl_path, "wb") as f: f.write(template_file.getbuffer())
            with open(raw_path, "wb") as f: f.write(raw_file.getbuffer())

            config = {
                "font": cfg_font, "paragraph": cfg_paragraph, "heading": cfg_heading,
                "table": cfg_table, "page_setup": cfg_page_setup, "preserve_char": cfg_preserve_char
            }

            try:
                update_progress(5, "Đang phân tích định dạng thị giác từ File mẫu...")
                engine = VisualTemplateEngine(tpl_path)
                if not engine.analyze()[0]:
                    st.error("Lỗi: Không thể phân tích File Mẫu.")
                else:
                    formatter = WordFormatterPipeline(raw_path, engine, out_path, config, update_progress)
                    success, msg = formatter.execute()
                    
                    if success:
                        st.success("🎉 Xử lý thành công! Nhấn nút bên dưới để tải về.")
                        with open(out_path, "rb") as f:
                            st.download_button(
                                label="⬇️ Tải xuống File Kết Quả",
                                data=f,
                                file_name=output_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    else:
                        st.error(f"Đã xảy ra lỗi trong quá trình xử lý: {msg}")
            except Exception as e:
                st.error(f"Lỗi hệ thống: {str(e)}")
